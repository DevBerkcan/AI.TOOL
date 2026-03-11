"""Document parser — extracts text from PDF, DOCX, HTML, Markdown."""

from __future__ import annotations

from dataclasses import dataclass

import structlog

logger = structlog.get_logger()


@dataclass
class ParsedDocument:
    """Output of document parsing."""

    title: str
    content: str
    metadata: dict
    source_type: str


async def parse_document(
    content: bytes,
    filename: str,
    mime_type: str,
    source_type: str = "pdf_upload",
) -> ParsedDocument:
    """Parse a document into structured text based on its MIME type."""

    if mime_type == "application/pdf" or filename.lower().endswith(".pdf"):
        return await _parse_pdf(content, filename, source_type)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or filename.lower().endswith(".docx"):
        return await _parse_docx(content, filename, source_type)
    elif mime_type == "text/html" or filename.lower().endswith(".html"):
        return _parse_html(content, filename, source_type)
    elif mime_type == "text/markdown" or filename.lower().endswith(".md"):
        return _parse_markdown(content, filename, source_type)
    elif mime_type == "text/plain" or filename.lower().endswith(".txt"):
        return ParsedDocument(
            title=filename,
            content=content.decode("utf-8", errors="replace"),
            metadata={"filename": filename},
            source_type=source_type,
        )
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type}")


async def _parse_pdf(content: bytes, filename: str, source_type: str) -> ParsedDocument:
    """Parse PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()

    full_text = "\n\n".join(text_parts).strip()

    if not full_text:
        logger.warning("PDF has no extractable text, may need OCR", filename=filename)

    return ParsedDocument(
        title=filename.rsplit(".", 1)[0],
        content=full_text,
        metadata={"filename": filename, "pages": len(text_parts)},
        source_type=source_type,
    )


async def _parse_docx(content: bytes, filename: str, source_type: str) -> ParsedDocument:
    """Parse DOCX using python-docx."""
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    return ParsedDocument(
        title=filename.rsplit(".", 1)[0],
        content=full_text,
        metadata={"filename": filename, "paragraphs": len(paragraphs)},
        source_type=source_type,
    )


def _parse_html(content: bytes, filename: str, source_type: str) -> ParsedDocument:
    """Parse HTML using BeautifulSoup."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")

    # Remove script, style, nav, footer elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    title = soup.title.string if soup.title else filename
    text = soup.get_text(separator="\n", strip=True)

    return ParsedDocument(
        title=str(title),
        content=text,
        metadata={"filename": filename},
        source_type=source_type,
    )


def _parse_markdown(content: bytes, filename: str, source_type: str) -> ParsedDocument:
    """Parse Markdown (kept as-is, it's already text)."""
    text = content.decode("utf-8", errors="replace")

    # Extract title from first heading
    title = filename
    for line in text.split("\n"):
        if line.startswith("# "):
            title = line[2:].strip()
            break

    return ParsedDocument(
        title=title,
        content=text,
        metadata={"filename": filename},
        source_type=source_type,
    )
